"""Parse uploaded files (PDF, DOCX, TXT) to plain text."""

from __future__ import annotations

import io
import logging

logger = logging.getLogger("kvp.parser")


def parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:  # noqa: BLE001
            logger.warning("pdf page %d failed: %s", i, e)
    return "\n\n".join(p.strip() for p in parts if p.strip())


def parse_docx(data: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Tables too — playbooks often use them
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def parse_txt(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def parse(filename: str, mime_type: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf") or mime_type == "application/pdf":
        return parse_pdf(data)
    if lower.endswith(".docx") or mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        return parse_docx(data)
    if lower.endswith((".txt", ".md", ".markdown")) or mime_type.startswith("text/"):
        return parse_txt(data)
    raise ValueError(f"Unsupported file type: {filename} ({mime_type})")
